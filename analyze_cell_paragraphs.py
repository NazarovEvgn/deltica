import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('docs/docx-templates/Приложение Ж (этикетки).docx')
table = doc.tables[0]

# Смотрим строки с контентом этикеток (1, 3, 5, 7)
for row_idx in [1, 3, 5, 7]:
    print(f"\n=== ROW {row_idx} (Label Content) ===")
    row = table.rows[row_idx]
    for cell_idx in [0, 2]:  # Колонки СИ и ИО
        cell = row.cells[cell_idx]
        print(f"\n  Cell [{row_idx},{cell_idx}]:")
        for para in cell.paragraphs:
            if para.text.strip():
                print(f"    {para.text}")
