import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document('docs/docx-templates/Приложение Ж (этикетки).docx')

print(f"=== TEMPLATE STRUCTURE ===")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print()

print("=== PARAGRAPHS (first 20) ===")
for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n=== TABLES ===")
for table_idx, table in enumerate(doc.tables):
    print(f"\nTable {table_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row_idx, row in enumerate(table.rows[:5]):  # First 5 rows
        cells_text = [cell.text.strip() for cell in row.cells]
        print(f"  Row {row_idx}: {cells_text}")
