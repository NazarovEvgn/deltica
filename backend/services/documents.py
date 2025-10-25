# deltica/backend/services/documents.py

from typing import Optional, List
from datetime import datetime
from pathlib import Path
from docxtpl import DocxTemplate
from sqlalchemy.orm import Session
from backend.app import models
from copy import deepcopy


class DocumentService:
    """Сервис для генерации документов из шаблонов"""

    def __init__(self, db: Session):
        self.db = db
        self.templates_dir = Path("docs/docx-templates")
        self.output_dir = Path("backend/generated_documents")
        self.output_dir.mkdir(exist_ok=True)

    def _get_equipment_full_data(self, equipment_id: int) -> Optional[dict]:
        """Получить полные данные оборудования для заполнения шаблонов"""
        equipment = self.db.query(models.Equipment).filter(
            models.Equipment.id == equipment_id
        ).first()

        if not equipment:
            return None

        # Получить связанные данные
        verification = self.db.query(models.Verification).filter(
            models.Verification.equipment_id == equipment_id
        ).first()

        responsibility = self.db.query(models.Responsibility).filter(
            models.Responsibility.equipment_id == equipment_id
        ).first()

        # Маппинг подразделений для отображения (должен совпадать с departmentOptions в EquipmentModal)
        department_map = {
            'gruppa_sm': 'Группа СМ',
            'gtl': 'ГТЛ',
            'lbr': 'ЛБР',
            'ltr': 'ЛТР',
            'lhaiei': 'ЛХАиЭИ',
            'ogmk': 'ОГМК',
            'oii': 'ОИИ',
            'smtsik': 'СМТСиК',
            'soii': 'СОИИ',
            'to': 'ТО',
            'ts': 'ТС',
            'es': 'ЭС',
            'ooops': 'ОООПС'  # Добавлено недостающее
        }

        # Форматирование дат
        verification_date = verification.verification_date.strftime('%d.%m.%Y') if verification and verification.verification_date else ''
        verification_due = verification.verification_due.strftime('%d.%m.%Y') if verification and verification.verification_due else ''
        department = department_map.get(responsibility.department, '') if responsibility else ''

        return {
            'equipment_name': equipment.equipment_name or '',
            'equipment_model': equipment.equipment_model or '',
            'factory_number': equipment.factory_number or '',
            'inventory_number': equipment.inventory_number or '',
            'verification_date': verification_date,
            'verification_due': verification_due,
            'department': department
        }

    def generate_label(self, equipment_id: int) -> Optional[str]:
        """
        Генерировать этикетку для оборудования
        Возвращает путь к сгенерированному файлу или None при ошибке
        """
        # Получить данные оборудования
        data = self._get_equipment_full_data(equipment_id)
        if not data:
            return None

        # Загрузить шаблон
        template_path = self.templates_dir / "template_label.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        template = DocxTemplate(template_path)

        # Заполнить шаблон
        template.render(data)

        # Сохранить результат
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"label_{equipment_id}_{timestamp}.docx"
        output_path = self.output_dir / output_filename

        template.save(str(output_path))

        return str(output_path)

    def generate_labels_batch(self, equipment_ids: List[int]) -> Optional[str]:
        """
        Генерировать пакет этикеток для нескольких единиц оборудования
        Использует простой шаблон, этикетки идут одна под другой без разрывов страниц
        Возвращает путь к сгенерированному файлу или None при ошибке
        """
        if not equipment_ids:
            return None

        # Получить данные для всех единиц оборудования
        equipments_data = []
        for equipment_id in equipment_ids:
            data = self._get_equipment_full_data(equipment_id)
            if data:  # Добавляем только найденное оборудование
                equipments_data.append(data)

        if not equipments_data:
            return None

        # Загрузить шаблон одной этикетки
        template_path = self.templates_dir / "template_label.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        from docx import Document
        from docx.shared import Cm

        # Генерируем первую этикетку - она станет основой документа
        template = DocxTemplate(template_path)
        template.render(equipments_data[0])

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"labels_{len(equipments_data)}_items_{timestamp}.docx"
        output_path = self.output_dir / output_filename

        template.save(str(output_path))

        # Если есть ещё этикетки, добавляем их
        if len(equipments_data) > 1:
            result_doc = Document(str(output_path))

            # Устанавливаем минимальные поля
            section = result_doc.sections[0]
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

            # Для остальных единиц оборудования генерируем этикетки
            for idx in range(1, len(equipments_data)):
                equipment_data = equipments_data[idx]

                # Загружаем шаблон и заполняем данными
                template = DocxTemplate(template_path)
                template.render(equipment_data)

                # Сохраняем во временный файл
                temp_path = self.output_dir / f"temp_label_{idx}.docx"
                template.save(str(temp_path))

                # Читаем заполненный шаблон
                temp_doc = Document(str(temp_path))

                # Копируем таблицу из шаблона
                if temp_doc.tables:
                    source_table = temp_doc.tables[0]
                    new_table_element = deepcopy(source_table._element)
                    result_doc.element.body.append(new_table_element)

                    # Принудительно устанавливаем выравнивание по левому краю
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn

                    # Находим элемент tblPr (table properties)
                    tblPr = new_table_element.find(qn('w:tblPr'))
                    if tblPr is not None:
                        # Удаляем существующее выравнивание если есть
                        jc = tblPr.find(qn('w:jc'))
                        if jc is not None:
                            tblPr.remove(jc)

                        # Добавляем выравнивание по левому краю
                        jc = OxmlElement('w:jc')
                        jc.set(qn('w:val'), 'left')
                        tblPr.append(jc)

                # Удаляем временный файл
                temp_path.unlink()

            # Сохраняем результат
            result_doc.save(str(output_path))

        return str(output_path)

    def generate_conservation_act(self, equipment_ids: List[int]) -> Optional[str]:
        """
        Генерировать акт консервации для нескольких единиц оборудования
        Возвращает путь к сгенерированному файлу или None при ошибке
        """
        if not equipment_ids:
            return None

        # Получить данные для всех единиц оборудования
        equipments_data = []
        for equipment_id in equipment_ids:
            data = self._get_equipment_full_data(equipment_id)
            if data:  # Добавляем только найденное оборудование
                equipments_data.append(data)

        if not equipments_data:
            return None

        # Загрузить шаблон акта консервации
        template_path = self.templates_dir / "template_storage.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        from docx import Document

        # Генерируем первую запись
        template = DocxTemplate(template_path)
        template.render(equipments_data[0])

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"conservation_act_{len(equipments_data)}_items_{timestamp}.docx"
        output_path = self.output_dir / output_filename

        template.save(str(output_path))

        # Если больше одной единицы оборудования, добавляем остальные строки
        if len(equipments_data) > 1:
            doc = Document(str(output_path))

            # Находим таблицу с данными (вторая таблица, индекс 1)
            data_table = doc.tables[1]

            # Для каждой дополнительной единицы оборудования добавляем строку
            for idx in range(1, len(equipments_data)):
                equipment = equipments_data[idx]
                row = data_table.add_row()

                # Порядковый номер с точкой
                row.cells[0].text = str(idx + 1) + "."

                # Данные оборудования
                row.cells[1].text = (
                    f"{equipment['equipment_name']}, "
                    f"зав. № {equipment['factory_number']}, "
                    f"инв. № {equipment['inventory_number']}"
                )

                # Применяем форматирование как в первой строке
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            from docx.shared import Pt
                            run.font.size = Pt(12)

            # Сохраняем финальный документ
            doc.save(str(output_path))

        return str(output_path)
