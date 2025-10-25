import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document('test_batch_labels_output.docx')
print(f'Всего таблиц: {len(doc.tables)}\n')

for idx, table in enumerate(doc.tables[:3]):
    print(f'=== Таблица {idx+1} (этикетка) ===')
    print(f'Строка 4, колонка 3: "{table.rows[4].cells[3].text.strip()}"')
    print(f'Строка 5, колонка 3: "{table.rows[5].cells[3].text.strip()}"')
    print()
