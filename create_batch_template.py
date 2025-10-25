import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from copy import deepcopy

# Загружаем оригинальный шаблон
doc = Document('docs/docx-templates/template_label.docx')
table = doc.tables[0]

# Создаём новый документ для пакетного шаблона
new_doc = Document()

# Добавляем Jinja2 цикл в начало (как параграф перед таблицей)
# Важно: docxtpl работает с XML-тегами, поэтому добавим маркер цикла
loop_start = new_doc.add_paragraph('{% for equipment in equipments %}')
loop_start.runs[0].font.size = Pt(1)  # Делаем текст очень маленьким (почти невидимым)

# Копируем таблицу из оригинала
new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
new_table.style = table.style

# Копируем содержимое ячеек
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        new_cell = new_table.rows[i].cells[j]
        # Копируем текст
        for paragraph in cell.paragraphs:
            if paragraph.text:
                new_para = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                new_para.text = paragraph.text
                # Копируем форматирование
                for run in paragraph.runs:
                    if run.text:
                        new_run = new_para.runs[0] if new_para.runs else new_para.add_run(run.text)
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                        new_run.font.bold = run.font.bold

        # Копируем ширину и границы
        new_cell.width = cell.width

# Добавляем разрыв страницы после таблицы (кроме последней итерации)
page_break = new_doc.add_paragraph('{% if not loop.last %}')
page_break.runs[0].font.size = Pt(1)
page_break_run = new_doc.add_paragraph()
page_break_run.add_run().add_break()  # Page break
loop_end_if = new_doc.add_paragraph('{% endif %}')
loop_end_if.runs[0].font.size = Pt(1)

# Закрываем цикл
loop_end = new_doc.add_paragraph('{% endfor %}')
loop_end.runs[0].font.size = Pt(1)

# Сохраняем новый шаблон
output_path = 'docs/docx-templates/template_labels_batch.docx'
new_doc.save(output_path)

print(f'✓ Пакетный шаблон создан: {output_path}')
print('\nВажно: Откройте шаблон в Word и проверьте, что:')
print('1. Jinja2 теги {% for %} и {% endfor %} присутствуют')
print('2. Таблица выглядит корректно')
print('3. После этого можно использовать для пакетной генерации')
