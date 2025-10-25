import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Загружаем шаблон
doc = Document('docs/docx-templates/template_label.docx')

# Добавляем параграф с началом цикла В САМОЕ НАЧАЛО документа
loop_start_para = doc.paragraphs[0].insert_paragraph_before('{% for equipment in equipments %}')
loop_start_para.runs[0].font.size = Pt(1)
loop_start_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет (невидимый)

# Добавляем условие для разрыва страницы и закрытие цикла В КОНЕЦ
# Добавляем условие перед разрывом страницы
if_para = doc.add_paragraph('{% if not loop.last %}')
if_para.runs[0].font.size = Pt(1)
if_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

# Добавляем сам разрыв страницы
page_break_para = doc.add_paragraph()
run = page_break_para.add_run()
# Добавляем XML элемент для разрыва страницы
run._element.append(OxmlElement('w:br'))
run._element[-1].set(qn('w:type'), 'page')

# Закрываем условие
endif_para = doc.add_paragraph('{% endif %}')
endif_para.runs[0].font.size = Pt(1)
endif_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

# Закрываем цикл
loop_end_para = doc.add_paragraph('{% endfor %}')
loop_end_para.runs[0].font.size = Pt(1)
loop_end_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

# Теперь заменяем все заполнители на equipment.XXX (в таблице)
table = doc.tables[0]
replacements = {
    '{{ equipment_name }}': '{{ equipment.equipment_name }}',
    '{{ equipment_model }}': '{{ equipment.equipment_model }}',
    '{{ factory_number }}': '{{ equipment.factory_number }}',
    '{{ inventory_number }}': '{{ equipment.inventory_number }}',
    '{{ verification_date }}': '{{ equipment.verification_date }}',
    '{{ verification_due }}': '{{ equipment.verification_due }}',
    '{{ department }}': '{{ equipment.department }}'
}

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

# Сохраняем как новый шаблон для пакетной генерации
output_path = 'docs/docx-templates/template_labels_batch.docx'
doc.save(output_path)

print(f'✓ Пакетный шаблон создан: {output_path}')
print('\nШаблон содержит:')
print('  - {% for equipment in equipments %} в начале')
print('  - {{ equipment.field_name }} вместо {{ field_name }}')
print('  - {% if not loop.last %} [page break] {% endif %}')
print('  - {% endfor %} в конце')
